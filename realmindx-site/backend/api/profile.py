from datetime import datetime, timedelta, timezone
import re
import secrets

from email_validator import EmailNotValidError, validate_email
from flask import Blueprint, Response, current_app, jsonify, request, send_file
from flask_login import current_user, login_required
from markupsafe import escape
from sqlalchemy.exc import IntegrityError
from werkzeug.security import check_password_hash, generate_password_hash

from ..audit import audit
from ..contacts import upsert_contact_safely
from ..checkout_details import checkout_detail_json, list_checkout_details, upsert_checkout_detail
from ..email_service import OutboundEmail, absolute_app_url, app_email_shell, send_admin_alert, send_email
from ..extensions import db, limiter
from ..location_data import canonical_delivery_locations, joined_location_ids, joined_location_names
from ..models import CheckoutDetail, ContactChangeToken, JobAlertPreference, TeacherPlacement, UploadedFile, User, UserProfile
from ..profile_completion import teacher_profile_completion
from ..serializers import user_json
from ..sms_service import normalise_phone, send_sms
from ..teacher_ids import ensure_application_id
from ..upload_utils import save_upload
from ..whatsapp_access import can_use_whatsapp_phone_verification
from ..whatsapp_service import (
    send_whatsapp_otp,
    whatsapp_business_number,
    whatsapp_challenge_phrase,
    whatsapp_challenge_url,
)

profile_bp = Blueprint("profile", __name__)

_LOCKED_STATUSES = frozenset({"submitted", "under_review", "rejected"})


def _queue_reverification(profile, reason, details=None):
    """Return an approved profile to review without changing its permanent Teacher ID."""
    if profile.profile_status != "verified":
        return False
    profile.profile_status = "submitted"
    profile.submitted_at = datetime.now(timezone.utc)
    profile.reviewed_at = None
    profile.reviewed_by_id = None
    profile.review_notes = None
    audit("teacher_profile_reverification_requested", "user", current_user.id, {
        "reason": reason,
        "teacher_id": current_user.teacher_id,
        **(details or {}),
    })
    return True


def _sync_editable_profile_status(profile):
    """Keep editable profile state aligned with the canonical completion score."""
    completion, missing = teacher_profile_completion(current_user)
    if profile.profile_status in {"incomplete", "revision_required"}:
        profile.profile_status = "complete" if completion >= 100 else "incomplete"
    return completion, missing


def _upload_url(uploaded_file):
    if not uploaded_file:
        return None
    return f"/uploads/{uploaded_file.visibility}/{uploaded_file.category}/{uploaded_file.stored_filename}"


def profile_json(profile):
    picture = db.session.get(UploadedFile, profile.profile_picture_file_id) if profile.profile_picture_file_id else None
    cv = db.session.get(UploadedFile, profile.cv_file_id) if profile.cv_file_id else None
    certificate = db.session.get(UploadedFile, profile.certificate_file_id) if profile.certificate_file_id else None
    placements = TeacherPlacement.query.filter_by(user_id=current_user.id).order_by(TeacherPlacement.accepted_at.desc()).all()
    completion, missing = teacher_profile_completion(current_user)
    return {
        "first_name": current_user.first_name,
        "last_name": current_user.last_name,
        "application_id": current_user.application_id,
        "teacher_id": current_user.teacher_id,
        "email": current_user.email,
        "phone": current_user.phone,
        "phone_verified": current_user.phone_verified,
        "whatsapp_phone_verification_allowed": can_use_whatsapp_phone_verification(current_user),
        "sex": current_user.sex,
        "age_range": current_user.age_range,
        "is_verified": current_user.is_verified,
        "location": profile.location,
        "teaching_subject": profile.teaching_subject,
        "preferred_level": profile.preferred_level,
        "preferred_employment_type": profile.preferred_employment_type,
        "available_from": profile.available_from,
        "curriculum_experience": profile.curriculum_experience,
        "preferred_locations": profile.preferred_locations,
        "preferred_location_ids": profile.preferred_location_ids,
        "bio": profile.bio,
        "profile_picture_file_id": profile.profile_picture_file_id,
        "profile_picture_url": _upload_url(picture),
        "cv_file_id": profile.cv_file_id,
        "cv_url": _upload_url(cv),
        "cv_filename": cv.original_filename if cv else None,
        "certificate_file_id": profile.certificate_file_id,
        "certificate_url": _upload_url(certificate),
        "certificate_filename": certificate.original_filename if certificate else None,
        "next_of_kin_name": profile.next_of_kin_name,
        "next_of_kin_phone": profile.next_of_kin_phone,
        "next_of_kin_relationship": profile.next_of_kin_relationship,
        "next_of_kin_email": profile.next_of_kin_email,
        "years_of_experience": profile.years_of_experience,
        "date_of_birth": profile.date_of_birth.isoformat() if profile.date_of_birth else None,
        "profile_status": profile.profile_status,
        "profile_completion": completion,
        "profile_missing_fields": missing,
        "submitted_at": profile.submitted_at.isoformat() if profile.submitted_at else None,
        "review_notes": profile.review_notes,
        "placements": [{
            "id": row.id, "school_name": row.school_name, "job_title": row.job_title,
            "status": row.status, "accepted_at": row.accepted_at.isoformat() if row.accepted_at else None,
            "started_at": row.started_at.isoformat() if row.started_at else None,
            "ended_at": row.ended_at.isoformat() if row.ended_at else None,
        } for row in placements],
    }


def _sync_profile_to_alert_preference(profile):
    """Bridge UserProfile → JobAlertPreference on first profile save.

    dispatch_job_alerts() in admin.py queries JobAlertPreference, NOT
    UserProfile.  Without this bridge, a user who sets their teaching subject
    and location in their profile never receives any job alerts, because no
    JobAlertPreference row exists for them.

    We only INSERT — never UPDATE — so a user who has already manually
    customised their alert preference keeps their customisation.
    """
    values = (
        profile.teaching_subject,
        profile.preferred_locations or profile.location,
        profile.preferred_level,
        profile.preferred_employment_type,
    )
    if not any(values):
        return
    pref = JobAlertPreference.query.filter_by(user_id=profile.user_id, is_default=True).first()
    if not pref:
        pref = JobAlertPreference(user_id=profile.user_id, is_default=True)
        db.session.add(pref)
    pref.subject = profile.teaching_subject
    pref.location = profile.preferred_locations or profile.location
    pref.location_ids = profile.preferred_location_ids
    pref.preferred_level = profile.preferred_level
    pref.curriculum = profile.curriculum_experience
    pref.employment_type = profile.preferred_employment_type
    pref.alert_by_email = True if pref.alert_by_email is None else pref.alert_by_email
    pref.frequency = pref.frequency or "instant"


def get_or_create_profile():
    profile = current_user.profile
    if not profile:
        profile = UserProfile(user_id=current_user.id)
        db.session.add(profile)
        db.session.flush()
    return profile


@profile_bp.get("/me/profile")
@login_required
def get_profile():
    return jsonify(profile=profile_json(get_or_create_profile()))


@profile_bp.put("/me/profile")
@login_required
def update_profile():
    payload = request.get_json(silent=True) or {}
    profile = get_or_create_profile()

    change_reason = str(payload.pop("change_reason", "") or "").strip()
    was_revision_required = profile.profile_status == "revision_required"
    was_verified = profile.profile_status == "verified"
    if was_verified and len(change_reason) < 8:
        return jsonify(error="Explain why you need this approved profile changed (at least 8 characters)."), 400
    before = {field: getattr(profile, field, None) for field in (
        "location", "teaching_subject", "preferred_level", "preferred_employment_type",
        "available_from", "curriculum_experience", "bio", "next_of_kin_name",
        "next_of_kin_phone", "next_of_kin_relationship", "next_of_kin_email",
        "preferred_location_ids", "years_of_experience", "date_of_birth",
    )}
    identity_before = {"first_name": current_user.first_name, "last_name": current_user.last_name}

    if profile.profile_status in _LOCKED_STATUSES:
        return jsonify(error="Your profile is currently under review and cannot be edited."), 423

    if "sex" in payload:
        current_user.sex = str(payload.get("sex") or "").strip().lower() or None
    if "age_range" in payload:
        current_user.age_range = str(payload.get("age_range") or "").strip().lower() or None
    if "first_name" in payload:
        first_name = str(payload.get("first_name") or "").strip()
        if not first_name:
            return jsonify(error="First name is required."), 400
        current_user.first_name = first_name
    if "last_name" in payload:
        current_user.last_name = str(payload.get("last_name") or "").strip() or None
    for field in [
        "location",
        "teaching_subject",
        "preferred_level",
        "preferred_employment_type",
        "available_from",
        "curriculum_experience",
        "bio",
        "next_of_kin_name",
        "next_of_kin_phone",
        "next_of_kin_relationship",
        "next_of_kin_email",
    ]:
        if field in payload:
            setattr(profile, field, payload[field])
    if "preferred_location_ids" in payload:
        try:
            zones, location_ids = canonical_delivery_locations(payload.get("preferred_location_ids"))
        except ValueError as exc:
            return jsonify(error=str(exc)), 400
        profile.preferred_location_ids = joined_location_ids(location_ids)
        profile.preferred_locations = joined_location_names(zones)
    if "years_of_experience" in payload:
        # Integer column — the form sends '' for "not selected", which Postgres
        # rejects outright ("invalid input syntax for type integer: ''"),
        # crashing the whole request with an unhandled 500. Coerce like
        # date_of_birth below: blank/invalid → NULL, valid → parsed value.
        raw = payload["years_of_experience"]
        if raw == "" or raw is None:
            profile.years_of_experience = None
        else:
            try:
                profile.years_of_experience = int(raw)
            except (ValueError, TypeError):
                pass
    if "date_of_birth" in payload:
        import datetime as _dt
        raw = payload["date_of_birth"]
        if raw:
            try:
                profile.date_of_birth = _dt.date.fromisoformat(raw)
            except (ValueError, TypeError):
                pass
        else:
            profile.date_of_birth = None
    if "phone" in payload:
        new_phone = payload["phone"] or None
        if new_phone != current_user.phone:
            return jsonify(error="Phone changes require OTP verification in Account & Security."), 400
    _sync_profile_to_alert_preference(profile)

    changed_fields = [field for field, previous in before.items() if getattr(profile, field, None) != previous]
    changed_fields.extend(field for field, previous in identity_before.items() if getattr(current_user, field) != previous)
    reverification = bool(changed_fields) and _queue_reverification(profile, change_reason or "profile_information_changed", {"changed_fields": changed_fields})
    if not reverification:
        _sync_editable_profile_status(profile)

    audit("profile_updated", "user_profile", current_user.id, {
        "email": current_user.email,
        "change_reason": change_reason or None,
        "changed_fields": changed_fields,
        "reverification_requested": reverification,
    })
    db.session.commit()

    if was_revision_required and profile.profile_status == "complete":
        try:
            admin_url = absolute_app_url("/admin/dashboard")
            send_admin_alert(
                subject=f"Teacher revision resubmitted: {current_user.application_id or current_user.email}",
                html=app_email_shell(
                    "Teacher resubmitted after revision",
                    (
                        f"<p><strong>Teacher:</strong> {escape(current_user.full_name or current_user.first_name or 'Teacher')}</p>"
                        f"<p><strong>Application ID:</strong> {escape(current_user.application_id or 'N/A')}</p>"
                        f"<p><strong>Email:</strong> {escape(current_user.email)}</p>"
                        "<p>The teacher has updated the profile after a revision was requested. "
                        "The profile is now complete and ready to be submitted for re-review.</p>"
                    ),
                    cta_label="Open Teacher Review Queue",
                    cta_url=admin_url,
                    eyebrow="RealMindX Internal: Teacher Resubmission",
                    preheader="A teacher resubmitted after revision.",
                ),
                text=(
                    f"Teacher revision resubmitted: {current_user.full_name or current_user.email}; "
                    f"application ID {current_user.application_id or 'N/A'}. Open {admin_url}"
                ),
                reply_to=current_user.email,
                template_name="teacher_revision_resubmitted_admin_alert",
            )
        except Exception:
            current_app.logger.exception("Revision resubmission alert failed for user %s", current_user.id)

    return jsonify(profile=profile_json(profile))


def _clean_email(value):
    try:
        return validate_email(value or "", check_deliverability=False).normalized.lower()
    except EmailNotValidError as exc:
        raise ValueError(str(exc)) from exc


def _mask_destination(field, value):
    if field == "email":
        name, domain = value.split("@", 1)
        visible = name[:2] if len(name) > 2 else name[:1]
        return f"{visible}{'*' * max(2, len(name) - len(visible))}@{domain}"
    return f"*** *** {value[-4:]}"


def _send_contact_change_code(field, target, code, channel="sms"):
    if field == "phone":
        if channel == "whatsapp":
            return send_whatsapp_otp(
                target,
                code,
                purpose="security",
                recipient_user_id=current_user.id,
                template_name="contact_change_phone_otp",
            )
        return send_sms(
            target,
            f"Your RealMindX verification code is {code}. It expires in 15 minutes. "
            "Do not share this code.",
            purpose="security",
            recipient_user_id=current_user.id,
            template_name="contact_change_phone_otp",
        )
    first_name = current_user.first_name or "there"
    body = (
        f"<p>Hello {escape(first_name)},</p>"
        "<p>Use the code below to verify your new email address for your RealMindX account.</p>"
        '<div style="text-align:center;margin:28px 0;">'
        '<div style="display:inline-block;background:#f5f8fc;border:2px dashed #c8d5e8;'
        'border-radius:12px;padding:18px 36px;">'
        f'<p style="margin:0;font-size:38px;font-weight:900;letter-spacing:.22em;color:#143670;">{escape(code)}</p>'
        '</div></div>'
        "<p style='font-size:13px;color:#6b80a0;'>This code expires in 15 minutes. "
        "If you did not request this change, keep your current account details and contact RealMindX.</p>"
    )
    return send_email(
        OutboundEmail(
            to=target,
            subject=f"Verify your RealMindX email change: {code}",
            html=app_email_shell(
                "Verify your new email",
                body,
                eyebrow="RealMindX Account Security",
                preheader=f"Your verification code is {code}. It expires in 15 minutes.",
            ),
        ),
        purpose="security",
        recipient_user_id=current_user.id,
        template_name="contact_change_email_otp",
    )


@profile_bp.put("/me/account")
@login_required
def update_account():
    payload = request.get_json(silent=True) or {}
    first_name = (payload.get("first_name") or "").strip()
    last_name = (payload.get("last_name") or "").strip()
    if not first_name:
        return jsonify(error="First name is required."), 400

    profile = current_user.profile
    if profile and profile.profile_status in _LOCKED_STATUSES:
        return jsonify(error="Identity details cannot be changed while your profile is locked."), 423

    change_reason = str(payload.get("change_reason") or "").strip()
    if profile and profile.profile_status == "verified" and len(change_reason) < 8:
        return jsonify(error="Explain why you need your approved name changed (at least 8 characters)."), 400
    previous = {"first_name": current_user.first_name, "last_name": current_user.last_name}
    current_user.first_name = first_name
    current_user.last_name = last_name or None
    names_changed = previous != {"first_name": current_user.first_name, "last_name": current_user.last_name}
    reverification = names_changed and _queue_reverification(profile, change_reason or "identity_name_changed", {
        "previous": previous,
        "current": {"first_name": current_user.first_name, "last_name": current_user.last_name},
    }) if profile else False
    audit("account_name_updated", "user", current_user.id, {"email": current_user.email, "change_reason": change_reason or None, "previous": previous, "reverification_requested": reverification})
    db.session.commit()
    return jsonify(user=user_json(current_user))


@profile_bp.get("/me/checkout-details")
@login_required
def get_checkout_details():
    return jsonify(items=list_checkout_details(current_user))


@profile_bp.post("/me/checkout-details")
@login_required
def save_checkout_details():
    payload = request.get_json(silent=True) or {}
    try:
        detail = upsert_checkout_detail(current_user.id, payload)
    except ValueError as exc:
        return jsonify(error=str(exc)), 400
    audit("checkout_details_saved", "checkout_detail", detail.id, {"user_id": current_user.id})
    db.session.commit()
    return jsonify(detail=checkout_detail_json(detail)), 201


@profile_bp.put("/me/checkout-details/<int:detail_id>")
@login_required
def update_checkout_details(detail_id):
    detail = CheckoutDetail.query.filter_by(id=detail_id, user_id=current_user.id).first_or_404()
    payload = request.get_json(silent=True) or {}
    merged = {
        "label": payload.get("label", detail.label),
        "customer_name": payload.get("customer_name", detail.customer_name),
        "email": payload.get("email", detail.email),
        "phone": payload.get("phone", detail.phone),
        "delivery_zone_id": payload.get("delivery_zone_id", detail.delivery_zone_id),
        "delivery_zone_name": payload.get("delivery_zone_name", detail.delivery_zone_name),
        "address": payload.get("address", detail.address),
        "city": payload.get("city", detail.city),
        "region": payload.get("region", detail.region),
        "is_default": payload.get("is_default", detail.is_default),
    }
    try:
        next_detail = upsert_checkout_detail(current_user.id, merged)
    except ValueError as exc:
        return jsonify(error=str(exc)), 400
    if next_detail.id != detail.id:
        db.session.delete(detail)
    audit("checkout_details_updated", "checkout_detail", next_detail.id, {"user_id": current_user.id})
    db.session.commit()
    return jsonify(detail=checkout_detail_json(next_detail))


@profile_bp.delete("/me/checkout-details/<int:detail_id>")
@login_required
def delete_checkout_details(detail_id):
    detail = CheckoutDetail.query.filter_by(id=detail_id, user_id=current_user.id).first_or_404()
    db.session.delete(detail)
    audit("checkout_details_deleted", "checkout_detail", detail_id, {"user_id": current_user.id})
    db.session.commit()
    return jsonify(message="Saved checkout details deleted.")


@profile_bp.post("/me/contact-change/request")
@login_required
@limiter.limit("6/hour")
def request_contact_change():
    payload = request.get_json(silent=True) or {}
    field = str(payload.get("field") or "").strip().lower()
    raw_value = str(payload.get("value") or "").strip()
    channel = str(payload.get("channel") or "sms").strip().lower()
    if field not in {"email", "phone"}:
        return jsonify(error="Choose email or phone verification."), 400
    if field == "phone" and channel not in {"sms", "whatsapp"}:
        return jsonify(error="Choose SMS or WhatsApp verification."), 400
    if field == "phone" and channel == "whatsapp":
        if not current_app.config.get("WHATSAPP_PHONE_VERIFICATION_ENABLED", False):
            return jsonify(error="WhatsApp verification is temporarily unavailable. Please use SMS for now."), 400
        if not can_use_whatsapp_phone_verification(current_user):
            return jsonify(error="WhatsApp verification is not available for this account right now. Please use SMS for now."), 400
    if field == "email":
        channel = "email"
    delivery_channel = "email"

    if field == "email":
        try:
            target = _clean_email(raw_value)
        except ValueError as exc:
            return jsonify(error=str(exc)), 400
        if target == current_user.email:
            return jsonify(error="That is already your account email."), 400
        existing = User.query.filter(User.email == target, User.id != current_user.id).first()
        if existing:
            return jsonify(error="That email is already connected to another account."), 409
    else:
        target = normalise_phone(raw_value)
        if not target:
            return jsonify(error="Enter a valid Ghana phone number."), 400
        if target == current_user.phone and current_user.phone_verified:
            return jsonify(error="That phone number is already verified."), 400
        if channel == "whatsapp" and current_app.config.get("WHATSAPP_INBOUND_CHALLENGE_ENABLED", True):
            delivery_channel = "whatsapp_inbound"
        elif channel == "whatsapp":
            delivery_channel = "whatsapp_template"
        else:
            delivery_channel = "sms"

    now = datetime.now(timezone.utc)
    window_start = now - timedelta(hours=1)
    recent_challenges = ContactChangeToken.query.filter(
        ContactChangeToken.user_id == current_user.id,
        ContactChangeToken.field == field,
        ContactChangeToken.created_at >= window_start,
    ).order_by(ContactChangeToken.created_at.desc()).all()
    if recent_challenges:
        latest_created_at = recent_challenges[0].created_at
        if latest_created_at.tzinfo is None:
            latest_created_at = latest_created_at.replace(tzinfo=timezone.utc)
        cooldown_seconds = 45 + (30 * (len(recent_challenges) - 1))
        elapsed_seconds = (now - latest_created_at).total_seconds()
        if elapsed_seconds < cooldown_seconds:
            retry_after = max(1, int(cooldown_seconds - elapsed_seconds + 0.999))
            return jsonify(
                error=f"Please wait {retry_after} seconds before requesting another code.",
                retry_after_seconds=retry_after,
            ), 429

    if delivery_channel == "whatsapp_inbound":
        ContactChangeToken.query.filter(
            ContactChangeToken.field == "phone",
            ContactChangeToken.delivery_channel == "whatsapp_inbound",
            ContactChangeToken.target_value == target,
            ContactChangeToken.used_at.is_(None),
            ContactChangeToken.expires_at < now,
        ).update({"status": "expired", "active_lock_key": None}, synchronize_session=False)

        active_phone_challenges = ContactChangeToken.query.filter(
            ContactChangeToken.field == "phone",
            ContactChangeToken.delivery_channel == "whatsapp_inbound",
            ContactChangeToken.target_value == target,
            ContactChangeToken.status == "pending",
            ContactChangeToken.used_at.is_(None),
            ContactChangeToken.expires_at >= now,
        ).all()
        if any(challenge.user_id != current_user.id for challenge in active_phone_challenges):
            return jsonify(
                error=(
                    "A WhatsApp verification for this number is already in progress. "
                    "Please wait a few minutes and try again."
                )
            ), 409
        for active_challenge in active_phone_challenges:
            active_challenge.used_at = now
            active_challenge.status = "cancelled"
            active_challenge.active_lock_key = None

    ContactChangeToken.query.filter_by(
        user_id=current_user.id,
        field=field,
        used_at=None,
    ).update({"used_at": now, "status": "cancelled", "active_lock_key": None})
    code = f"{secrets.randbelow(1_000_000):06d}"
    challenge = ContactChangeToken(
        user_id=current_user.id,
        field=field,
        target_value=target,
        delivery_channel=delivery_channel,
        status="pending",
        active_lock_key=target if delivery_channel == "whatsapp_inbound" else None,
        token_hash=generate_password_hash(code),
        expires_at=now + timedelta(minutes=15),
    )
    db.session.add(challenge)
    try:
        db.session.flush()
    except IntegrityError:
        db.session.rollback()
        if delivery_channel == "whatsapp_inbound":
            return jsonify(
                error=(
                    "A WhatsApp verification for this number is already in progress. "
                    "Please wait a few minutes and try again."
                )
            ), 409
        raise
    if delivery_channel == "whatsapp_inbound":
        delivered = True
        mocked = False
    else:
        result = _send_contact_change_code(field, target, code, channel)
        delivered = result.status in ("queued", "accepted", "sent", "delivered")
        mocked = result.status == "mocked" and current_app.config.get("ENV") != "production"
    if not delivered and not mocked:
        db.session.rollback()
        current_app.logger.warning("Could not deliver %s verification code for user %s", field, current_user.id)
        channel_label = "WhatsApp" if channel == "whatsapp" else field
        return jsonify(error=f"Could not send the verification code by {channel_label}. Please try again."), 502
    audit(
        "contact_change_requested",
        "user",
        current_user.id,
        {"field": field, "channel": channel, "destination": _mask_destination(field, target)},
    )
    db.session.commit()
    masked = _mask_destination(field, target)
    delivery_label = "WhatsApp" if channel == "whatsapp" else channel.upper() if channel == "sms" else channel.title()
    if delivery_channel == "whatsapp_inbound":
        phrase = whatsapp_challenge_phrase()
        business_number = whatsapp_business_number()
        return jsonify(
            challenge_id=challenge.id,
            field=field,
            channel=channel,
            delivery_channel=delivery_channel,
            verification_mode="whatsapp_inbound",
            destination=masked,
            target_phone=target,
            challenge_phrase=phrase,
            whatsapp_number=business_number,
            whatsapp_url=whatsapp_challenge_url(phrase),
            expires_in_seconds=900,
            next_request_in_seconds=45 + (30 * len(recent_challenges)),
            manual_entry_allowed=False,
            message=f"Open WhatsApp and send the prefilled verification message to {business_number} from {masked}.",
        )
    if mocked:
        return jsonify(
            challenge_id=challenge.id,
            field=field,
            channel=channel,
            delivery_channel=delivery_channel,
            delivery_status="mocked",
            destination=masked,
            expires_in_seconds=900,
            next_request_in_seconds=45 + (30 * len(recent_challenges)),
            message="Verification request recorded in local mock mode. No real message was sent.",
        ), 202
    return jsonify(
        challenge_id=challenge.id,
        field=field,
        channel=channel,
        delivery_channel=delivery_channel,
        delivery_status="accepted",
        destination=masked,
        expires_in_seconds=900,
        next_request_in_seconds=45 + (30 * len(recent_challenges)),
        message=f"Verification code sent by {delivery_label} to {masked}.",
    )


@profile_bp.get("/me/contact-change/<int:challenge_id>/status")
@login_required
def contact_change_status(challenge_id):
    challenge = ContactChangeToken.query.filter_by(
        id=challenge_id,
        user_id=current_user.id,
    ).first()
    if not challenge:
        return jsonify(error="Verification request not found."), 404

    expires_at = challenge.expires_at
    if expires_at and expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    expired = not expires_at or expires_at < datetime.now(timezone.utc)
    verified = bool(
        challenge.status == "verified"
        and challenge.used_at
        and challenge.field == "phone"
        and current_user.phone == challenge.target_value
        and current_user.phone_verified
    )
    wrong_number = (
        not verified
        and not expired
        and challenge.delivery_channel == "whatsapp_inbound"
        and challenge.last_whatsapp_attempt_status == "wrong_number"
    )
    status = (
        "verified" if verified else
        "wrong_number" if wrong_number else
        "expired" if expired or challenge.status == "expired" else
        "cancelled" if challenge.status == "cancelled" else
        "pending"
    )
    wrong_number_message = None
    if wrong_number:
        attempted = _mask_destination("phone", challenge.last_whatsapp_attempt_from or "")
        expected = _mask_destination("phone", challenge.target_value)
        wrong_number_message = (
            f"The challenge was sent from a different WhatsApp number ({attempted}), but you are verifying {expected}. "
            "Open WhatsApp with the account for the number you entered, or change the number below."
        )
    return jsonify(
        challenge_id=challenge.id,
        field=challenge.field,
        delivery_channel=challenge.delivery_channel,
        destination=_mask_destination(challenge.field, challenge.target_value),
        status=status,
        verified=verified,
        wrong_number=wrong_number,
        wrong_message=False,
        last_attempt_from=_mask_destination("phone", challenge.last_whatsapp_attempt_from or "") if challenge.last_whatsapp_attempt_from else None,
        last_attempt_status=challenge.last_whatsapp_attempt_status,
        last_attempt_at=challenge.last_whatsapp_attempt_at.isoformat() if challenge.last_whatsapp_attempt_at else None,
        expires_at=challenge.expires_at.isoformat() if challenge.expires_at else None,
        used_at=challenge.used_at.isoformat() if challenge.used_at else None,
        message=(
            "Phone verified."
            if verified else wrong_number_message
            if wrong_number else "Challenge expired. Send a fresh one."
            if status == "expired" else "This verification request was replaced. Start a fresh one if you still need to verify this number."
            if status == "cancelled" else "Waiting for the WhatsApp verification message. Send the prefilled message without changing it."
        ),
    )


@profile_bp.post("/me/contact-change/verify")
@login_required
@limiter.limit("12/hour")
def verify_contact_change():
    payload = request.get_json(silent=True) or {}
    try:
        challenge_id = int(payload.get("challenge_id"))
    except (TypeError, ValueError):
        return jsonify(error="Verification request not found. Send a fresh code."), 400
    otp = re.sub(r"\D", "", str(payload.get("otp") or ""))
    if len(otp) != 6:
        return jsonify(error="Enter the 6 digit verification code."), 400

    challenge = ContactChangeToken.query.filter_by(
        id=challenge_id,
        user_id=current_user.id,
        used_at=None,
    ).first()
    if not challenge:
        return jsonify(error="Verification code expired. Send a fresh code."), 400
    if challenge.delivery_channel == "whatsapp_inbound":
        return jsonify(error="Send this WhatsApp challenge from the phone number being verified. It cannot be typed into the site."), 400
    expires_at = challenge.expires_at
    if expires_at and expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if not expires_at or expires_at < datetime.now(timezone.utc):
        challenge.status = "expired"
        challenge.active_lock_key = None
        db.session.commit()
        return jsonify(error="Verification code expired. Send a fresh code."), 400
    if not check_password_hash(challenge.token_hash, otp):
        return jsonify(error="Verification code is incorrect."), 400

    now = datetime.now(timezone.utc)
    if challenge.field == "email":
        existing = User.query.filter(
            User.email == challenge.target_value,
            User.id != current_user.id,
        ).first()
        if existing:
            return jsonify(error="That email is already connected to another account."), 409
        current_user.email = challenge.target_value
        current_user.is_verified = True
    elif challenge.field == "phone":
        current_user.phone = challenge.target_value
        current_user.phone_verified = True
        current_user.phone_verified_at = now
    else:
        return jsonify(error="Unsupported verification request."), 400

    challenge.used_at = now
    challenge.status = "verified"
    challenge.verified_at = now
    challenge.active_lock_key = None
    if current_user.teacher_service_enabled:
        upsert_contact_safely(
            current_user.email,
            full_name=current_user.full_name,
            phone=current_user.phone,
            source="teacher",
            source_record_id=current_user.id,
            metadata={"application_id": current_user.application_id},
            activity_at=now,
            logger=current_app.logger,
        )
    ContactChangeToken.query.filter(
        ContactChangeToken.user_id == current_user.id,
        ContactChangeToken.field == challenge.field,
        ContactChangeToken.id != challenge.id,
        ContactChangeToken.used_at.is_(None),
    ).update({"used_at": now, "status": "cancelled", "active_lock_key": None})
    audit("contact_change_verified", "user", current_user.id, {"field": challenge.field})
    db.session.commit()
    return jsonify(
        message=f"{challenge.field.title()} updated and verified.",
        user=user_json(current_user),
    )


@profile_bp.post("/me/uploads")
@login_required
def upload_user_file():
    file = request.files.get("file")
    kind = request.form.get("kind") or "documents"
    field_map = {
        "profile_picture": ("images", "profile_picture_file_id", "public"),
        "cv": ("documents", "cv_file_id", "protected"),
        "certificate": ("documents", "certificate_file_id", "protected"),
        "document": ("documents", None, "protected"),
    }
    category, profile_field, visibility = field_map.get(kind, ("documents", None, "protected"))
    profile = get_or_create_profile()

    change_reason = str(request.form.get("change_reason") or "").strip()
    if profile.profile_status == "verified" and profile_field and len(change_reason) < 8:
        return jsonify(error="Explain why you need to replace this approved profile item (at least 8 characters)."), 400

    if profile_field in ("cv_file_id", "certificate_file_id") and profile.profile_status in _LOCKED_STATUSES:
        return jsonify(error="Your profile is under review and documents cannot be replaced."), 423

    try:
        uploaded = save_upload(file, category=category, owner_id=current_user.id, visibility=visibility)
    except ValueError as exc:
        return jsonify(error=str(exc)), 400
    db.session.flush()
    if profile_field:
        previous_file_id = getattr(profile, profile_field)
        setattr(profile, profile_field, uploaded.id)
        reverification = _queue_reverification(profile, change_reason or f"{kind}_replaced", {
            "previous_file_id": previous_file_id,
            "new_file_id": uploaded.id,
        })
        if not reverification:
            _sync_editable_profile_status(profile)
    audit("file_uploaded", "uploaded_file", uploaded.id, {
        "kind": kind, "filename": uploaded.original_filename, "category": uploaded.category,
        "change_reason": change_reason or None,
    })
    db.session.commit()
    return jsonify(
        file_id=uploaded.id,
        original_filename=uploaded.original_filename,
        category=uploaded.category,
        url=_upload_url(uploaded),
        profile=profile_json(profile),
    ), 201


def _can_access_teacher_file(uploaded):
    if uploaded.owner_id == current_user.id:
        return True
    role_name = current_user.role.name if current_user.role else None
    return role_name == "admin" or (role_name == "staff" and current_user.has_permission("teachers.view"))


def _docx_preview_html(path, filename):
    from docx import Document

    document = Document(path)
    body = []
    for paragraph in document.paragraphs:
        text = str(escape(paragraph.text or ""))
        if text:
            style = (paragraph.style.name or "").lower() if paragraph.style else ""
            tag = "h1" if "title" in style or "heading 1" in style else "h2" if "heading 2" in style else "p"
            body.append(f"<{tag}>{text}</{tag}>")
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = "".join(f"<td>{escape(cell.text or '')}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        body.append(f"<table>{''.join(rows)}</table>")
    return f"""<!doctype html><html><head><meta charset=\"utf-8\"><title>{escape(filename)}</title><style>
body{{font:15px/1.6 Arial,sans-serif;color:#172554;max-width:900px;margin:0 auto;padding:48px;background:#fff}}h1,h2{{color:#082d6b}}table{{border-collapse:collapse;width:100%;margin:18px 0}}td{{border:1px solid #dbe3ef;padding:8px;vertical-align:top}}p{{white-space:pre-wrap}}
</style></head><body>{''.join(body) or '<p>This document has no previewable text.</p>'}</body></html>"""


@profile_bp.get("/files/<int:file_id>/preview")
@login_required
def preview_uploaded_file(file_id):
    uploaded = db.get_or_404(UploadedFile, file_id)
    if not _can_access_teacher_file(uploaded):
        return jsonify(error="You do not have permission to view this file."), 403
    path = uploaded.storage_path
    extension = (uploaded.original_filename.rsplit(".", 1)[-1] if "." in uploaded.original_filename else "").lower()
    audit("teacher_document_viewed", "uploaded_file", uploaded.id, {"owner_id": uploaded.owner_id, "filename": uploaded.original_filename})
    db.session.commit()
    if extension == "pdf":
        response = send_file(path, mimetype="application/pdf", as_attachment=False, download_name=uploaded.original_filename)
    elif extension == "docx":
        try:
            response = Response(_docx_preview_html(path, uploaded.original_filename), mimetype="text/html")
        except Exception:
            current_app.logger.exception("Could not render DOCX preview for file %s", uploaded.id)
            return Response("<h2>Preview unavailable</h2><p>This Word document could not be rendered. Use Download to open it in Word.</p>", status=422, mimetype="text/html")
    else:
        return Response("<h2>Preview unavailable</h2><p>This file type cannot be displayed in the browser. Use Download to open it.</p>", status=415, mimetype="text/html")
    response.headers["Cache-Control"] = "private, no-store"
    response.headers["X-Content-Type-Options"] = "nosniff"
    # Document previews are embedded only by the authenticated RealMindX UI.
    # Keep cross-origin framing blocked while permitting that same-origin iframe.
    response.headers["X-Frame-Options"] = "SAMEORIGIN"
    return response


@profile_bp.get("/files/<int:file_id>/download")
@login_required
def download_uploaded_file(file_id):
    uploaded = db.get_or_404(UploadedFile, file_id)
    if not _can_access_teacher_file(uploaded):
        return jsonify(error="You do not have permission to download this file."), 403
    audit("teacher_document_downloaded", "uploaded_file", uploaded.id, {"owner_id": uploaded.owner_id, "filename": uploaded.original_filename})
    db.session.commit()
    return send_file(uploaded.storage_path, as_attachment=True, download_name=uploaded.original_filename, mimetype=uploaded.mime_type or None)


def _send_submission_email(user, profile):
    first_name = user.first_name or "Teacher"
    dashboard_url = absolute_app_url("/portal?view=profile")
    body = (
        f"<p>Dear {escape(first_name)},</p>"
        "<p>Thank you for submitting your teacher application. Your profile and documents have been received.</p>"
        f"<p><strong>Application ID:</strong> {escape(user.application_id or 'N/A')}</p>"
        f"<p><strong>Submission date:</strong> {profile.submitted_at.strftime('%B %d, %Y') if profile.submitted_at else 'N/A'}</p>"
        "<p>Your application is now in the review queue. A RealMindX administrator will review your profile and documents.</p>"
        "<p>What happens next:</p>"
        "<ul>"
        "<li>An administrator will review your profile and documents</li>"
        "<li>If everything is in order, your profile will be verified and a permanent Teacher ID will be issued</li>"
        "<li>If changes are needed, we will notify you with details of what to update</li>"
        "</ul>"
        "<p>You do not need to take any further action at this time. You will receive an email once the review is complete or if corrections are requested.</p>"
    )
    try:
        send_email(
            OutboundEmail(
                to=user.email,
                subject="Your Teacher Application Has Been Submitted",
                html=app_email_shell(
                    "Application Submitted",
                    body,
                    cta_label="View Dashboard",
                    cta_url=dashboard_url,
                    eyebrow="RealMindX Teacher Verification",
                    preheader="Your teacher application has been submitted and is in the review queue.",
                ),
            ),
            purpose="transactional",
            recipient_user_id=user.id,
            template_name="teacher_profile_submitted",
        )
    except Exception as exc:
        current_app.logger.warning(
            "Submission email failed for user %s (error=%s)",
            user.id,
            type(exc).__name__,
        )


def _send_submission_admin_alert(user, profile):
    admin_url = absolute_app_url("/admin/dashboard")
    submitted_at = profile.submitted_at.strftime("%B %d, %Y at %H:%M UTC") if profile.submitted_at else "N/A"
    is_resubmit = bool(getattr(profile, "review_notes", None))
    subject_prefix = "Teacher application resubmitted" if is_resubmit else "Teacher application submitted"
    heading = "Teacher application resubmitted for review" if is_resubmit else "Teacher application ready for review"
    body_note = (
        "<p>The teacher has addressed the requested revisions and the profile is now waiting for re-review.</p>"
        if is_resubmit else
        "<p>The completed profile and documents are now waiting in the teacher review queue.</p>"
    )
    template = "teacher_profile_resubmitted_admin_alert" if is_resubmit else "teacher_profile_submitted_admin_alert"
    send_admin_alert(
        subject=f"{subject_prefix}: {user.application_id or user.email}",
        html=app_email_shell(
            heading,
            (
                f"<p><strong>Teacher:</strong> {escape(user.full_name or user.first_name or 'Teacher')}</p>"
                f"<p><strong>Application ID:</strong> {escape(user.application_id or 'N/A')}</p>"
                f"<p><strong>Email:</strong> {escape(user.email)}</p>"
                f"<p><strong>Submitted:</strong> {escape(submitted_at)}</p>"
                + body_note
            ),
            cta_label="Open Teacher Review Queue",
            cta_url=admin_url,
            eyebrow="RealMindX Internal: Teacher Review",
            preheader="A teacher application is ready for admin review.",
        ),
        text=(
            f"{heading}: {user.full_name or user.email}; "
            f"application ID {user.application_id or 'N/A'}. Open {admin_url}"
        ),
        reply_to=user.email,
        template_name=template,
    )


@profile_bp.post("/me/profile/submit")
@login_required
def submit_profile():
    profile = get_or_create_profile()

    if not current_user.teacher_service_enabled:
        return jsonify(error="Teacher profile submission is not available for this account."), 403

    previous_application_id = current_user.application_id
    ensure_application_id(current_user)
    if current_user.application_id != previous_application_id:
        # Persist the identity repair even if a later submission validation
        # returns early (for example, an already-submitted legacy profile).
        db.session.commit()

    # Lock the row so two concurrent requests cannot both pass the status check.
    locked = db.session.query(UserProfile).with_for_update().filter_by(id=profile.id).first()
    profile = locked or profile

    if profile.profile_status == "submitted":
        return jsonify(
            profile_status="submitted",
            submitted_at=profile.submitted_at.isoformat() if profile.submitted_at else None,
            message="Your profile has already been submitted for review."
        ), 200

    if profile.profile_status == "revision_required":
        return jsonify(error="Please update your profile and address the requested changes before submitting."), 400

    if profile.profile_status == "rejected":
        return jsonify(error="Your profile has been rejected and cannot be resubmitted."), 400

    if profile.profile_status == "verified":
        return jsonify(error="Your profile has already been verified."), 400

    if profile.profile_status not in ("complete", "incomplete"):
        return jsonify(error="Profile cannot be submitted in its current state."), 400

    completion, missing = teacher_profile_completion(current_user)
    if completion < 100:
        return jsonify(error="Your profile must be 100% complete before submitting."), 400

    if not profile.cv_file_id:
        return jsonify(error="A CV is required before submitting."), 400

    if not profile.certificate_file_id:
        return jsonify(error="A certificate is required before submitting."), 400

    profile.profile_status = "submitted"
    profile.submitted_at = datetime.now(timezone.utc)

    audit("teacher_profile_submitted", "user", current_user.id, {
        "application_id": current_user.application_id,
        "submitted_at": profile.submitted_at.isoformat(),
    })
    db.session.commit()

    _send_submission_email(current_user, profile)
    _send_submission_admin_alert(current_user, profile)

    return jsonify(
        profile_status="submitted",
        submitted_at=profile.submitted_at.isoformat(),
        message="Your profile has been submitted for review."
    ), 200
